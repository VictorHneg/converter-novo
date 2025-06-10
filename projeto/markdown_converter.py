import re
import docx
import argparse
import os
import sys
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import qn
import io

class MarkdownToDocxConverter:
    def extract_text_from_docx(self, docx_path):
        """Extrai texto de um arquivo .docx"""
        doc = docx.Document(docx_path)
        text = []
        for para in doc.paragraphs:
            text.append(para.text)
        return '\n'.join(text)

    def add_document_property(doc, name, value):
        """Adiciona uma propriedade personalizada ao documento"""
        prop = doc.core_properties
        try:
            setattr(prop, name, value)
        except AttributeError:
            # Se não for uma propriedade padrão, adiciona como propriedade personalizada
            custom_props = prop._element.xpath('./cp:coreProperties/cp:customProperties', 
                                              namespaces=prop._element.nsmap)
            if not custom_props:
                custom_props = OxmlElement('cp:customProperties')
                prop._element.xpath('./cp:coreProperties')[0].append(custom_props)
            
            custom_prop = OxmlElement('cp:customProperty')
            custom_prop.set(qn('cp:name'), name)
            custom_prop.set(qn('cp:value'), value)
            custom_props[0].append(custom_prop)

    def parse_markdown_to_docx(self, markdown_text, output_path, drive_id=None):
        doc = docx.Document()

        # Configura margens do documento
        for section in doc.sections:
            section.left_margin = Pt(72)
            section.right_margin = Pt(72)
            section.top_margin = Pt(72)
            section.bottom_margin = Pt(72)

        lines = markdown_text.split('\n')
        current_section = None

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Remove asteriscos duplicados do markdown
            line = re.sub(r'\*\*(.*?)\*\*', r'\1', line)

            # Título principal (primeira linha em maiúsculas longa)
            if line.isupper() and len(line) > 10:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                para.space_after = Pt(12)
                continue

            # Campos com dois pontos (ex: "NOME DO PARTICIPANTE:")
            if ':' in line and not line.startswith('-') and not line.startswith('•'):
                parts = line.split(':', 1)
                para = doc.add_paragraph()
                run = para.add_run(parts[0].strip() + ':')
                run.bold = True
                if len(parts) > 1 and parts[1].strip():
                    para.add_run(' ' + parts[1].strip())
                para.space_after = Pt(6)
                continue

            # Seções principais (ex: "FONOAUDIOLOGIA")
            if line.isupper() and not ':' in line:
                current_section = line
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.bold = True
                para.space_before = Pt(12)
                para.space_after = Pt(6)
                continue

            # Lista com marcadores
            if line.startswith('-') or line.startswith('•'):
                text = line.lstrip('-• ').strip()
                if not text:
                    continue  # pula marcadores vazios
                para = doc.add_paragraph()
                para.style = 'List Bullet'
                # Negrito para itens especiais
                if text.lower().startswith(('metas terapêuticas', 'objetivos terapêuticos', 'observações')):
                    parts = text.split(':', 1)
                    run = para.add_run(parts[0] + ':')
                    run.bold = True
                    if len(parts) > 1:
                        para.add_run(parts[1])
                else:
                    para.add_run(text)
                para.space_after = Pt(2)
                continue

            # Observações ou texto normal
            para = doc.add_paragraph()
            if current_section and 'OBSERVAÇÕES' in current_section.upper():
                para.paragraph_format.left_indent = Pt(20)
            para.add_run(line)
            para.space_after = Pt(2)

        # Salva em bytes para uso com Flask send_file
        doc_bytes = io.BytesIO()
        doc.save(doc_bytes)
        doc_bytes.seek(0)
        return doc_bytes

    def process_inline_formatting(paragraph, text):
        """Processa formatação inline como negrito e itálico"""
        # Encontra todos os padrões de formatação
        parts = []
        current_pos = 0
        
        # Padrão para negrito: **texto**
        bold_pattern = re.compile(r'\*\*(.*?)\*\*')
        # Padrão para itálico: *texto* ou _texto_
        italic_pattern = re.compile(r'(\*|_)(.*?)(\*|_)')
        
        # Mescla os padrões com posição para processamento
        formats = []
        
        # Encontra negrito
        for match in bold_pattern.finditer(text):
            formats.append((match.start(), match.end(), match.group(1), 'bold'))
        
        # Encontra itálico
        for match in italic_pattern.finditer(text):
            # Evita processar ** novamente (já tratado como negrito)
            if not (match.group(1) == '*' and match.group(3) == '*' and 
                    match.start() > 0 and text[match.start()-1] == '*' and 
                    match.end() < len(text) and text[match.end()] == '*'):
                formats.append((match.start(), match.end(), match.group(2), 'italic'))
        
        # Ordena por posição inicial
        formats.sort(key=lambda x: x[0])
        
        # Se não houver formatação, adicione o texto inteiro
        if not formats:
            paragraph.add_run(text)
            return
        
        # Processa o texto com as formatações
        last_end = 0
        for start, end, content, format_type in formats:
            # Adiciona texto não formatado antes do formato atual
            if start > last_end:
                paragraph.add_run(text[last_end:start])
            
            # Adiciona o texto formatado
            run = paragraph.add_run(content)
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'italic':
                run.italic = True
            
            last_end = end
        
        # Adiciona qualquer texto restante após o último formato
        if last_end < len(text):
            paragraph.add_run(text[last_end:])

    @staticmethod
    def show_usage_and_exit():
        """Mostra uma mensagem de uso amigável e sai"""
        print("\nConversor de Markdown para DOCX")
        print("===============================\n")
        print("Este programa converte arquivos .docx com texto markdown para documentos formatados.")
        print("\nUso:")
        print("  python markdown_to_docx.py arquivo_entrada.docx arquivo_saida.docx [--drive-id ID]\n")
        print("Exemplos:")
        print("  python markdown_to_docx.py documento.docx documento_formatado.docx")
        print("  python markdown_to_docx.py documento.docx documento_formatado.docx --drive-id 1abc123\n")
        print("Para assistência interativa, execute:")
        print("  converter_markdown.bat\n")
        sys.exit(1)

    @staticmethod
    def main():
        parser = argparse.ArgumentParser(description='Converte markdown em arquivos .docx para documentos Word formatados')
        parser.add_argument('input_file', help='Caminho para o arquivo .docx ou ID do Google Drive')
        parser.add_argument('output_file', help='Caminho para salvar o arquivo .docx formatado')
        parser.add_argument('--drive-id', '-d', help='ID opcional do Google Drive para incluir no documento')
        parser.add_argument('--google-doc', '-g', action='store_true', help='Indica que o input é um ID do Google Drive')

        if len(sys.argv) == 1:
            MarkdownToDocxConverter.show_usage_and_exit()
        
        args = parser.parse_args()
        converter = MarkdownToDocxConverter()
        
        try:
            if args.google_doc:
                from google_drive_integration import process_drive_file
                from google.oauth2.credentials import Credentials
                from googleapiclient.discovery import build
                from google_auth import get_google_drive_credentials
                
                try:
                    credentials = get_google_drive_credentials()
                    service = build('drive', 'v3', credentials=credentials)
                except Exception as e:
                    print(f"Erro na autenticação do Google Drive: {str(e)}")
                    return 1
                    
                docx_file, filename = process_drive_file(service, args.input_file)
                markdown_text = converter.extract_text_from_docx(docx_file)
            else:
                if not os.path.exists(args.input_file):
                    print(f"Erro: Arquivo de entrada '{args.input_file}' não encontrado.")
                    return 1
                markdown_text = converter.extract_text_from_docx(args.input_file)
            
            doc_bytes = converter.parse_markdown_to_docx(markdown_text, args.output_file, args.drive_id)
            # Salva o arquivo de saída
            with open(args.output_file, 'wb') as f:
                f.write(doc_bytes.read())
            
            print(f"\nConversão concluída com sucesso!")
            print(f"Documento formatado salvo em: {args.output_file}")
            if args.drive_id:
                print(f"ID do Google Drive adicionado: {args.drive_id}")
            return 0
        except Exception as e:
            print(f"\nErro durante a conversão: {str(e)}")
            return 1

if __name__ == "__main__":
    sys.exit(MarkdownToDocxConverter.main())
